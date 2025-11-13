import datetime

from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from domain_types.chapters import Chapter, Theme, Lesson, SelfWorkTheme, SelfWorkPracticeTheme, PracticeLesson
from domain_types.organization import Organization, Worker, Department
from domain_types.subject import Subject, Room, Books, ExamQuestions
from view_types.doc_task_view_model import DocumentTaskView
from view_types.program_view_model import ProgramViewModel
from view_types.view_skill_type import SkillView, AimPersonalSkillsView, ProfessionalSkillsView, ResultsView


def remove_marked_rows(doc, marker="{delete}"):
    """
    Удаляет строки таблицы, содержащие указанный маркер
    """
    for table in doc.tables:
        # Собираем индексы строк для удаления (снизу вверх)
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if marker in cell.text:
                    rows_to_delete.append(i)
                    break

        # Удаляем строки в обратном порядке
        for i in sorted(rows_to_delete, reverse=True):
            table._tbl.remove(table.rows[i]._tr)

def find_table_index_after_phrase(doc, phrase="паспорт фонда"):
    found_phrase = False

    for i, block in enumerate(doc.element.body):
        # Проверяем параграф
        if block.tag.endswith('p'):
            paragraph = Paragraph(block, doc)
            if phrase.lower() in paragraph.text.lower():
                found_phrase = True

        # Проверяем таблицу
        elif block.tag.endswith('tbl'):
            if found_phrase:
                # Это первая таблица после фразы
                # Находим её индекс среди таблиц документа
                for table_index, table in enumerate(doc.tables):
                    if table._element == block:
                        return table_index

    return None

def merge_identical_cells_in_column(doc, phrase, column_index=2):
    table = doc.tables[find_table_index_after_phrase(doc, phrase)]

    start_cell = None
    last_text = None

    for i, row in enumerate(table.rows):
        cell = row.cells[column_index]
        text = cell.text.strip()

        # если значение совпадает с предыдущим — продолжаем группу
        if text == last_text:
            cell.text = ''
            continue
        else:
            # если была группа одинаковых ячеек — объединяем
            if start_cell and i - start_row > 1:
                end_cell = table.rows[i - 1].cells[column_index]
                start_cell.merge(end_cell)

            # начинаем новую группу
            start_cell = cell
            start_row = i
            last_text = text

    # объединяем последнюю группу, если нужно
    if start_cell and start_row < len(table.rows) - 1:
        end_cell = table.rows[-1].cells[column_index]
        if end_cell.text.strip() == last_text:
            start_cell.merge(end_cell)

def generate_document_with_deletion(template_path, output_path, context):
    """
    Генерирует документ и удаляет помеченные строки
    """
    doc = DocxTemplate(template_path)
    doc.render(context)

    remove_marked_rows(doc, marker="{delete}")

    merge_identical_cells_in_column(doc, 'Паспорт оценочных материалов')

    doc.save(output_path)
    print(f"Документ создан: {output_path}")





# Пример использования с простыми переменными
if __name__ == "__main__":
    docTask = DocumentTaskView(
        '{delete}',
        Organization(
        'Автономная некоммерческая организация '
        'профессионального образования «Колледж экономики, '
        'страхового дела и информационных технологий КЭСИ» '
        '(АНО ПО «Колледж КЭСИ»)',
        'Москва'),
        [Worker('Корнеев',
                'Антон',
                'Дмитриевич'),
         Worker('Комышан',
                'Пётр',
                'Иванович')],
        Department('информационных систем и программирования',
                   Worker('Комышан', 'Пётр', 'Иванович')
                   ),
        2000,
        ProgramViewModel(
            '09.02.07',
            'Информационные системы и программирование',
            1547,
            '26 декабря 2016',
            'Программист',
            2,
        ),
        Subject(
            'ОП.09',
            'Стандартизация, сертификация и метрология',
            True,
            False,
            2,
            0,
            Room(
                'Общепрофессиональных дисциплин и профессиональных модулей',
                [
                    'Стол ученический двухместный, нерегулируемый',
                    'Стул ученический на ножках',
                    'Стол учителя',
                    'Стул учителя'
                ]
            ),
            [
                "12 Моноблоков Lenovo 27' Разрешение экрана: 1920х1080 ЦП: AMD Ryzen 5 5500u ОЗУ: ddr4 8gb Ssd: 512gb"
            ],
            Books(
                ['Овчинникова, Е. А. Информационная безопасность : учебное пособие для СПО / Е. А. Овчинникова. — Саратов : Профобразование, 2024. — 166 c. — ISBN 978-5-4488-1872-1. — Текст : электронный // Электронный ресурс цифровой образовательной среды СПО PROFобразование : [сайт]. — URL: https://profspo.ru/books/139028 (дата обращения: 01.04.2025)'],
                [
                    'Фомин, Д. В. Информационная безопасность : учебно-методическое пособие для студентов заочной формы обучения направления подготовки 38.03.05 «Бизнес-информатика» / Д. В. Фомин. — Саратов : Вузовское образование, 2018. — 125 c. — ISBN 978-5-4487-0299-0. — Текст : электронный // Электронный ресурс цифровой образовательной среды СПО PROFобразование : [сайт]. — URL: https://profspo.ru/books/77318 (дата обращения: 01.04.2025)',
                    'Мирзакулова, Ш. А. Защита информации. Архитектура безопасности OSI : учебное пособие для ТиПО / Ш. А. Мирзакулова. — Алматы, Саратов : EDP Hub (Идипи Хаб), Профобразование, 2024. — 126 c. — ISBN 978-5-4488-2228-5. — Текст : электронный // Электронный ресурс цифровой образовательной среды СПО PROFобразование : [сайт]. — URL: https://profspo.ru/books/142546 (дата обращения: 01.04.2025)'
                ],
                ['Мега источник', 'Вау, копец какой источник']
            ),
            ExamQuestions(
                2,
                [
                    'Информационная безопасность, её основные задачи',
                    'Классы безопасности информационных систем',
                    'Основные стандарты безопасности операционных систем'
                ],
                [
                    'Создайте в Windows 10 нового пользователя "audit_user" без прав администратора и установите для него квоту диска 500 МБ.',
                ]
            )
        ),
        AimPersonalSkillsView(
            [
                SkillView('ОК 1',
                      'Понимать сущность и социальную значимость своей будущей профессии, проявлять к ней устойчивый интерес.'),
                SkillView('ОК 2',
                      'Организовывать собственную деятельность, выбирать типовые методы и способы выполнения профессиональных задач, оценивать их выполнение и качество.'),
                SkillView('ОК 3',
                      'Принимать решения в стандартных и нестандартных ситуациях и нести за них ответственность.')
            ],
            ProfessionalSkillsView(
                'ВД 2',
                'Администрирование баз данных',
                [
                    SkillView('ПК 2.1', 'Выявлять проблемы, возникающие в процессе эксплуатации баз данных'),
                    SkillView('ПК 2.3',
                          'Проводить аудит систем безопасности баз данных с использованием регламентов по защите информации.')
                ]
            ),
            [
                SkillView('ЛР 1', 'Быть крутым челом')
            ],
            [
                SkillView('ЦО 1', 'Быть целевым челом')
            ],
        ),
        ResultsView(
            ['анализировать угрозы информационной безопасности',
             'применять базовые методы защиты информации',
             'использовать антивирусные и межсетевые средства защиты'],
            ['основные виды угроз информационной безопасности',
             'методы и средства защиты информации',
             'правовые основы информационной безопасности',
             'принципы организации защиты информации в операционных системах']
        ),

        [
            Chapter(
                'Концепция информационной безопасности',
                1,
                [
                    Theme(
                        'Сущность и понятие информационной безопасности',
                        [
                            Lesson(
                                'Понятие информационной безопасности. Характеристика составляющих информационной безопасности. Источники и содержание угроз в информационной сфере.',
                                1
                            ),
                            Lesson(
                                'Состояние информационной безопасности Российской Федерации и основные задачи по ее обеспечению.',
                                1
                            )
                        ],
                       ['ОК 1', 'ОК 2', 'КекЛол']
                    ),

                    SelfWorkTheme(
                        'Сущность и понятие информационной безопасности',
                        [
                            Lesson(
                                'Понятие информационной безопасности. Характеристика составляющих информационной безопасности. Источники и содержание угроз в информационной сфере.',
                                9
                            ),
                            Lesson(
                                'Состояние информационной безопасности Российской Федерации и основные задачи по ее обеспечению.',
                                1
                            )
                        ],
                        ['ОК 1', 'ОК 2', 'КекЛол'],
                        [
                            Lesson(
                                'Национальные интересы Российской Федерации в информационной сфере. Влияние процессов информатизации общества на составляющие национальной безопасности и их содержание.',
                                17
                            )
                        ],
                    ),

                    SelfWorkPracticeTheme(
                        'Сущность и понятие информационной безопасности',
                        [
                            Lesson(
                                'Понятие информационной безопасности. Характеристика составляющих информационной безопасности. Источники и содержание угроз в информационной сфере.',
                                1
                            ),
                            Lesson(
                                'Состояние информационной безопасности Российской Федерации и основные задачи по ее обеспечению.',
                                1
                            )
                        ],
                        ['ОК 1', 'ОК 2', 'КекЛол'],
                        [
                            Lesson(
                                'Национальные интересы Российской Федерации в информационной сфере. Влияние процессов информатизации общества на составляющие национальной безопасности и их содержание.',
                                2
                            )
                        ],
                        [
                            PracticeLesson(
                                'Осуществление загрузки ОС с различных источников',
                                2,
                                1
                            ),
                            PracticeLesson(
                                'Осуществление загрузки ОС с различных источников',
                                2,
                                2
                            )

                        ]
                    ),

                ]
            )
        ]
    )

    generate_document_with_deletion(
        "templates/cesi_template_op.docx",
        f"output/{docTask.subject.code}_{docTask.year}_{datetime.datetime.now().strftime('%H-%M-%S')}.docx",
        docTask.__dict__
    )
